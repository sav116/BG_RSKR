# ----------------------
# -*- coding: utf-8 -*-
# Script by Artem.S.
# ----------------------

import docx, os, shutil, datetime, openpyxl
from settings.config import ppo_d, spo_d


def parsing_log(file):

    #This function returns a description of the error
    #param file: 'C:\Users\Artem\PycharmProjects\RSKR\PPO\10.19.88.2\MongoDB_findstr\MongoDB-20210415-0000(7_days)_Error.txt'
    #return: 'ERROR:  syntax error at or near "as" at character 35 ... '

    errors = []
    with open(file, encoding='utf-8') as f:
        for line in f:
            if 'error' in line.lower():
                start = line.find('ERROR')
            elif 'fatal' in line.lower():
                start = line.find('FATAL')
            elif '[Warning]' in line:
                start = line.find('[Warning]')
            errors.append(line[start:])
    return ''.join(errors)


def find_service_name(str):

    #This function returns service name from the string:
    #param str: '\\PPO\\10.19.88.2\\MongoDB_findstr\\'
    #return: 'MongoDB' or ''

    index1 = None
    index2 = None
    for i in range(len(str)):
        if index1 != None and index2 != None:
            return ' - ' + str[index1 + 2:index2].title() + '.'
        elif str[i].isdigit() and str[i + 1] == '\\':
            index1 = i
        elif str[i] == '_':
            index2 = i
    return ''


def get_discription_iis_code(sheet, errors, ip):
    result = ''
    is_description = []
    for error in errors:
        for row in range(2, sheet.max_row + 1):
            error_code = str(sheet.cell(row=row, column=1).value)
            if error in error_code:
                description = sheet.cell(row=row, column=2).value
                result += f"\n{description}\n"
                is_description.append(error)

    is_not_description = [i for i in errors if i not in is_description]
    if is_not_description:
        print(f"Не найдены описания error code: {is_not_description} для {ip}")
    return result


def get_discription_id(sheet, errors, ip):
    result = ''
    is_description = []
    for error in errors:
        for row in range(2, sheet.max_row + 1):
            id = str(sheet.cell(row=row, column=1).value)
            if ',' in id:
                id = [i.strip() for i in id.split(',')]
            if error in id:
                description = sheet.cell(row=row, column=2).value
                is_description.append(error)
                if description not in result:
                    if isinstance(id, list):
                        id = ', '.join(id)
                    result += f"\nEvent ID {id} - {description}\n"
                    break

    is_not_description = [i for i in errors if i not in is_description]
    if is_not_description:
        print(f"Не найдены описания для EventID: {is_not_description} для {ip}")
    return result


def get_iis_errors(current_path, ip, sheet_codes):
    iis_log_file = current_path + spo_d[ip]['iis'] + os.listdir(current_path + spo_d[ip]['iis'])[
        0]  # absolute path to log file iis
    temp_errors_list = []  # list for unique errors code
    with open(iis_log_file, encoding='UTF-16 LE') as f:
        for line in f:
            status_code = line.split(' ')[-4]
            if status_code not in temp_errors_list and status_code != '200':
                temp_errors_list.append(status_code)
    if temp_errors_list == []:
        return 'Ошибок и предупреждений не обнаружено'
    else:
        result = get_discription_iis_code(sheet_codes, temp_errors_list, ip)
    return result


def get_win_err_id(current_path, ip, sheet_id):
    win_app_log_file = current_path + spo_d[ip]['win_app'] + os.listdir(current_path + spo_d[ip]['win_app'])[
        0]  # absolute path to log file win_app
    win_sys_log_file = current_path + spo_d[ip]['win_sys'] + os.listdir(current_path + spo_d[ip]['win_sys'])[
        0]  # absolute path to log file win_sys
    temp_id_list = []  # list for unique id
    with open(win_app_log_file, encoding='UTF-16 LE') as f:
        for line in f:
            if 'EventID' in line:
                id = ''.join([i for i in line if i.isdigit()])
                if id not in temp_id_list:
                    temp_id_list.append(id)

    with open(win_sys_log_file, encoding='UTF-16 LE') as f:
        for line in f:
            if 'EventID' in line:
                id = ''.join([i for i in line if i.isdigit()])
                if id not in temp_id_list:
                    temp_id_list.append(id)

    if temp_id_list == []:
        return 'Ошибок и предупреждений не обнаружено'
    else:
        result = get_discription_id(sheet_id, temp_id_list, ip)
        return result

def get_apache_errors(current_path, ip):
    apache_log_file =  current_path + spo_d[ip]['apache_tomcat'] + os.listdir(current_path + spo_d[ip]['apache_tomcat'])[0]
    if os.path.getsize(apache_log_file) == 0:
        return "Ошибок и предупреждений не обнаружено"

def get_RedHat_errors(current_path, ip):
    RedHat_log_file =  current_path + spo_d[ip]['red_hat'] + os.listdir(current_path + spo_d[ip]['red_hat'])[0]
    errors=''
    if os.path.getsize(RedHat_log_file) == 0:
        return "Ошибок и предупреждений не обнаружено"
    else:
        with open(RedHat_log_file, encoding='UTF-8') as f:
            for line in f:
                errors+=line+'\n'
    return errors

def run_spo():
    now = datetime.datetime.now()  # object datetime.datetime 2021-06-10 11:14:10.166054
    current_path = os.getcwd()  # absolute path to scripts
    new_table = f"SPO_{now.day}_{now.month}_{now.year}.xlsx"  # name for the new file
    shutil.copyfile(current_path + '\\settings\\table_default.xlsx',
                    current_path + "\\" + new_table)  # copy template table file from folder 'settings' like new SPO table
    tab_spo = openpyxl.load_workbook(filename=new_table)
    tab_id_bd = openpyxl.load_workbook(filename=f"settings\\event_id.xlsx")
    tab_iis_bd = openpyxl.load_workbook(filename=f"settings\\iis_errors.xlsx")
    sheet_tab_spo = tab_spo.active  # open active book in which we will write the results
    sheet_tab_id_bd = tab_id_bd.active  # open active book of table with description by event id
    sheet_tab_iis_bd = tab_iis_bd.active  # open active book of table with iis errors code and description
    # iterate over all lines
    ip = None
    for row in range(2, 22):
        if '10.19.88.' in str(sheet_tab_spo.cell(row=row, column=3).value):
            ip = sheet_tab_spo.cell(row=row, column=3).value.strip()
        component = sheet_tab_spo.cell(row=row, column=5).value.lower()  # cell value is equal to the lower case
        if 'window' in component:
            sheet_tab_spo.cell(row=row, column=6).value = get_win_err_id(current_path, ip, sheet_tab_id_bd)
        elif 'iis' in component:
            sheet_tab_spo.cell(row=row, column=6).value = get_iis_errors(current_path, ip, sheet_tab_iis_bd)
        elif 'apache' in component:
            sheet_tab_spo.cell(row=row, column=6).value = get_apache_errors(current_path, ip)
        elif 'red hat' in component:
            sheet_tab_spo.cell(row=row, column=6).value = get_RedHat_errors(current_path, ip)
        elif 'httpd' in component:
            sheet_tab_spo.cell(row=row, column=6).value = get_RedHat_errors(current_path, ip)
    tab_spo.save(new_table)


def run_ppo():
    current_path = os.getcwd()
    doc_ppo = docx.Document()
    doc_ppo.add_heading('ППО', 0)  # add tittle for type document
    for ip in ppo_d:
        doc_ppo.add_heading(ip + find_service_name(ppo_d[ip]), 1)
        if 'нет' in ppo_d[ip]:
            p = doc_ppo.add_paragraph(ppo_d[ip] + '\n')
            p.add_run().bold = True
        elif 'ошибки' in ppo_d[ip]:
            p = doc_ppo.add_paragraph(ppo_d[ip] + '\n')
            p.add_run().bold = True
        else:
            try:
                abs_path_to_err_log = current_path + ppo_d[ip] + os.listdir(current_path + ppo_d[ip])[0]
                if os.path.getsize(abs_path_to_err_log) > 0:  # select not empty files
                    p = doc_ppo.add_paragraph(parsing_log(abs_path_to_err_log) + '\n')
                    p.add_run().bold = True
                elif os.path.getsize(abs_path_to_err_log) == 0:
                    p = doc_ppo.add_paragraph('Ошибок не обнаружено' + '\n')
                    p.add_run().bold = True
            except:
                p = doc_ppo.add_paragraph('Лог файл не найден' + '\n')
                p.add_run().bold = True
    return doc_ppo

def send_reports_to_telegram():
    pass

if __name__ == '__main__':
    # run PowerShell script and waiting when it stopped
    now = datetime.datetime.now()
    doc_ppo = run_ppo()
    doc_ppo.save(f'PPO_{now.day}_{now.month}_{now.year}.doc')
    doc_spo = run_spo()
    send_reports_to_telegram()

